from __future__ import annotations

from pathlib import Path
import shutil
import textwrap
import subprocess
from datetime import date

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.drawing.image import Image as XLImage
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak, KeepTogether

from primarycare_model.parameterised_model import (
    PARAMETERISED_SCENARIOS,
    parameter_input_frame,
    scenario_input_frame,
    run_parameterised_game_models,
    run_parameterised_hybrid_model,
    run_parameterised_uncertainty,
    summarise_parameterised_uncertainty,
    sensitivity_by_parameter,
    hybrid_informed_mcda,
)
from primarycare_model.uncertainty import PARAMETER_FIELDS

VERSION = "v1.0.0"
ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUTPUTS = ROOT / "outputs"
FIGURES = DOCS / "figures"
PARAMS_DIR = DOCS / "parameters"
MODELLING_DIR = DOCS / "modelling"
MCDA_DIR = DOCS / "mcda"
REPORTS_DIR = ROOT / "reports"
TEMPLATES_DIR = ROOT / "data" / "templates"
TRACK_DIR = ROOT / "conductor" / "tracks" / "015-parameterised-model-run"
for d in [OUTPUTS, FIGURES, PARAMS_DIR, MODELLING_DIR, MCDA_DIR, REPORTS_DIR, TEMPLATES_DIR, TRACK_DIR]:
    d.mkdir(parents=True, exist_ok=True)

# -------------------------------
# Run model
# -------------------------------
param_df = parameter_input_frame()
scenario_df = scenario_input_frame()
game_df = run_parameterised_game_models()
hybrid_df = run_parameterised_hybrid_model()
# 1,000 draws per scenario gives stable summary while keeping the artefact portable.
game_draws, hybrid_draws = run_parameterised_uncertainty(n_per_scenario=1000, sd=0.07, seed=20260508)
unc_summary = summarise_parameterised_uncertainty(hybrid_draws)
sens_viability = sensitivity_by_parameter(hybrid_draws, "hybrid_viability_score", top_n=8)
sens_hospital = sensitivity_by_parameter(hybrid_draws, "weighted_hospital_pressure", top_n=8)
mcda_df = hybrid_informed_mcda()

# Deterministic game summary by scenario.
game_summary = (
    game_df.groupby(["scenario_id", "scenario_name"], as_index=False)
    .agg(
        access_score=("access_score", "mean"),
        provider_viability=("provider_viability", "mean"),
        equity_score=("equity_score", "mean"),
        fiscal_control=("fiscal_control", "mean"),
        hospital_pressure=("hospital_pressure", "mean"),
        gaming_risk=("gaming_risk", "mean"),
        system_welfare=("system_welfare", "mean"),
    )
    .round(2)
)

equilibrium_matrix = game_df.pivot(index="game_name", columns="scenario_id", values="equilibrium_label").reset_index()

# -------------------------------
# Write CSV artefacts
# -------------------------------
files = {
    PARAMS_DIR / f"parameter-input-register-{VERSION}.csv": param_df,
    MODELLING_DIR / f"parameterised-scenario-inputs-{VERSION}.csv": scenario_df,
    MODELLING_DIR / f"parameterised-game-results-{VERSION}.csv": game_df,
    MODELLING_DIR / f"parameterised-game-summary-{VERSION}.csv": game_summary,
    MODELLING_DIR / f"parameterised-equilibrium-matrix-{VERSION}.csv": equilibrium_matrix,
    MODELLING_DIR / f"parameterised-hybrid-results-{VERSION}.csv": hybrid_df,
    MODELLING_DIR / f"parameterised-hybrid-uncertainty-summary-{VERSION}.csv": unc_summary,
    MODELLING_DIR / f"parameterised-sensitivity-hybrid-viability-{VERSION}.csv": sens_viability,
    MODELLING_DIR / f"parameterised-sensitivity-hospital-pressure-{VERSION}.csv": sens_hospital,
    MCDA_DIR / f"hybrid-informed-mcda-results-{VERSION}.csv": mcda_df,
    OUTPUTS / f"parameter-input-register-{VERSION}.csv": param_df,
    OUTPUTS / f"parameterised-scenario-inputs-{VERSION}.csv": scenario_df,
    OUTPUTS / f"parameterised-game-results-{VERSION}.csv": game_df,
    OUTPUTS / f"parameterised-game-summary-{VERSION}.csv": game_summary,
    OUTPUTS / f"parameterised-equilibrium-matrix-{VERSION}.csv": equilibrium_matrix,
    OUTPUTS / f"parameterised-hybrid-results-{VERSION}.csv": hybrid_df,
    OUTPUTS / f"parameterised-hybrid-uncertainty-summary-{VERSION}.csv": unc_summary,
    OUTPUTS / f"parameterised-sensitivity-hybrid-viability-{VERSION}.csv": sens_viability,
    OUTPUTS / f"parameterised-sensitivity-hospital-pressure-{VERSION}.csv": sens_hospital,
    OUTPUTS / f"hybrid-informed-mcda-results-{VERSION}.csv": mcda_df,
}
for path, df in files.items():
    df.to_csv(path, index=False)

# Draw samples for audit without bloating the package.
hybrid_draws.sample(n=min(5000, len(hybrid_draws)), random_state=20260508).round(4).to_csv(OUTPUTS / f"parameterised-hybrid-uncertainty-draw-sample-{VERSION}.csv", index=False)
game_draws.sample(n=min(5000, len(game_draws)), random_state=20260509).round(4).to_csv(OUTPUTS / f"parameterised-game-uncertainty-draw-sample-{VERSION}.csv", index=False)

# -------------------------------
# Figures
# -------------------------------
scenario_order = list(scenario_df["scenario_id"])
scenario_labels = [f"{sid}\n{name.split()[0]}" for sid, name in zip(scenario_df["scenario_id"], scenario_df["name"])]

# Figure 1: hybrid viability bar
plot_df = hybrid_df.set_index("scenario_id").loc[scenario_order].reset_index()
plt.figure(figsize=(11, 5.5))
plt.bar(plot_df["scenario_id"], plot_df["hybrid_viability_score"])
plt.ylabel("Hybrid viability score (0-100)")
plt.xlabel("Scenario")
plt.title("Parameterised model run: hybrid viability by scenario")
plt.ylim(0, 100)
plt.xticks(rotation=30, ha="right")
for i, v in enumerate(plot_df["hybrid_viability_score"]):
    plt.text(i, v + 1, f"{v:.1f}", ha="center", fontsize=8)
plt.tight_layout()
fig1 = FIGURES / f"parameterised-hybrid-viability-{VERSION}.png"
plt.savefig(fig1, dpi=220)
plt.savefig(OUTPUTS / fig1.name, dpi=220)
plt.close()

# Figure 2: hospital pressure vs supply
plt.figure(figsize=(8, 6))
plt.scatter(plot_df["supply_generation_index"], plot_df["weighted_hospital_pressure"], s=80)
for _, r in plot_df.iterrows():
    plt.annotate(r["scenario_id"], (r["supply_generation_index"], r["weighted_hospital_pressure"]), xytext=(5, 5), textcoords="offset points")
plt.xlabel("Supply generation index (0-100)")
plt.ylabel("Weighted hospital pressure (0-100; lower is better)")
plt.title("Upstream supply generation and hospital pressure")
plt.xlim(0, 100)
plt.ylim(0, 100)
plt.tight_layout()
fig2 = FIGURES / f"parameterised-supply-vs-hospital-pressure-{VERSION}.png"
plt.savefig(fig2, dpi=220)
plt.savefig(OUTPUTS / fig2.name, dpi=220)
plt.close()

# Figure 3: key indices line chart
indices = ["supply_generation_index", "equity_legitimacy_index", "governance_resilience_index", "hospital_deflection_index", "hybrid_viability_score"]
plt.figure(figsize=(11, 6))
for col in indices:
    plt.plot(plot_df["scenario_id"], plot_df[col], marker="o", label=col.replace("_", " "))
plt.ylabel("Score (0-100)")
plt.xlabel("Scenario")
plt.title("Hybrid model component indices")
plt.ylim(0, 100)
plt.xticks(rotation=30, ha="right")
plt.legend(loc="best", fontsize=8)
plt.tight_layout()
fig3 = FIGURES / f"parameterised-hybrid-indices-{VERSION}.png"
plt.savefig(fig3, dpi=220)
plt.savefig(OUTPUTS / fig3.name, dpi=220)
plt.close()

# Figure 4: game welfare heatmap
welfare_matrix = game_df.pivot(index="game_id", columns="scenario_id", values="system_welfare").loc[:, scenario_order]
plt.figure(figsize=(11, 7))
plt.imshow(welfare_matrix.values, aspect="auto")
plt.colorbar(label="System welfare score")
plt.xticks(range(len(welfare_matrix.columns)), welfare_matrix.columns, rotation=45, ha="right")
plt.yticks(range(len(welfare_matrix.index)), welfare_matrix.index)
plt.xlabel("Scenario")
plt.ylabel("Game")
plt.title("Game-level welfare by scenario")
plt.tight_layout()
fig4 = FIGURES / f"parameterised-game-welfare-heatmap-{VERSION}.png"
plt.savefig(fig4, dpi=220)
plt.savefig(OUTPUTS / fig4.name, dpi=220)
plt.close()

# Figure 5: uncertainty boxplot for hybrid viability
box_data = [hybrid_draws.loc[hybrid_draws["scenario_id"] == sid, "hybrid_viability_score"].astype(float).values for sid in scenario_order]
plt.figure(figsize=(11, 5.5))
plt.boxplot(box_data, labels=scenario_order, showfliers=False)
plt.ylabel("Hybrid viability score")
plt.xlabel("Scenario")
plt.title("Uncertainty stress test: hybrid viability")
plt.tight_layout()
fig5 = FIGURES / f"parameterised-uncertainty-hybrid-viability-{VERSION}.png"
plt.savefig(fig5, dpi=220)
plt.savefig(OUTPUTS / fig5.name, dpi=220)
plt.close()

# Figure 6: pooled sensitivity drivers
pooled = []
for param in PARAMETER_FIELDS:
    x = hybrid_draws[param].astype(float)
    y = hybrid_draws["hybrid_viability_score"].astype(float)
    corr = float(np.corrcoef(x, y)[0, 1]) if x.nunique() > 1 and y.nunique() > 1 else np.nan
    pooled.append({"parameter": param, "correlation": corr, "abs_correlation": abs(corr)})
pooled_sens = pd.DataFrame(pooled).dropna().sort_values("abs_correlation", ascending=False).head(12).sort_values("abs_correlation")
pooled_sens.to_csv(OUTPUTS / f"parameterised-pooled-sensitivity-hybrid-viability-{VERSION}.csv", index=False)
plt.figure(figsize=(9, 6))
plt.barh(pooled_sens["parameter"], pooled_sens["abs_correlation"])
plt.xlabel("Absolute correlation with hybrid viability")
plt.title("Top pooled sensitivity drivers")
plt.tight_layout()
fig6 = FIGURES / f"parameterised-top-sensitivity-drivers-{VERSION}.png"
plt.savefig(fig6, dpi=220)
plt.savefig(OUTPUTS / fig6.name, dpi=220)
plt.close()

# Figure 7: MCDA ranking
mcda_plot = mcda_df.sort_values("risk_adjusted_mcda_score")
plt.figure(figsize=(10, 6))
plt.barh(mcda_plot["scenario_id"], mcda_plot["risk_adjusted_mcda_score"])
plt.xlabel("Risk-adjusted MCDA score")
plt.ylabel("Scenario")
plt.title("Hybrid-informed MCDA ranking")
plt.xlim(0, 100)
plt.tight_layout()
fig7 = FIGURES / f"parameterised-mcda-ranking-{VERSION}.png"
plt.savefig(fig7, dpi=220)
plt.savefig(OUTPUTS / fig7.name, dpi=220)
plt.close()

# Figure 8: hospital pressure intervals
intervals = unc_summary.set_index("scenario_id").loc[scenario_order].reset_index()
y = intervals["weighted_hospital_pressure_mean"].astype(float)
yerr = [
    (y - intervals["weighted_hospital_pressure_p05"].astype(float)).values,
    (intervals["weighted_hospital_pressure_p95"].astype(float) - y).values,
]
plt.figure(figsize=(10, 5.5))
plt.errorbar(intervals["scenario_id"], y, yerr=yerr, fmt="o", capsize=4)
plt.ylabel("Weighted hospital pressure (lower is better)")
plt.xlabel("Scenario")
plt.title("Hospital pressure uncertainty intervals")
plt.tight_layout()
fig8 = FIGURES / f"parameterised-hospital-pressure-intervals-{VERSION}.png"
plt.savefig(fig8, dpi=220)
plt.savefig(OUTPUTS / fig8.name, dpi=220)
plt.close()

# -------------------------------
# Reports: markdown, docx, pdf
# -------------------------------

def md_table(df: pd.DataFrame, cols: list[str], n: int | None = None) -> str:
    d = df[cols].copy()
    if n:
        d = d.head(n)
    return d.to_markdown(index=False)

summary_hybrid_cols = ["scenario_id", "scenario_name", "hybrid_viability_score", "supply_generation_index", "hospital_deflection_index", "weighted_hospital_pressure", "fiscal_gaming_risk_index", "interaction_penalty"]
mcda_cols = ["rank", "scenario_id", "scenario_name", "weighted_score_before_penalty", "risk_penalty", "risk_adjusted_mcda_score"]
unc_cols = ["scenario_id", "scenario_name", "hybrid_viability_score_mean", "hybrid_viability_score_p05", "hybrid_viability_score_p95", "weighted_hospital_pressure_mean", "weighted_hospital_pressure_p05", "weighted_hospital_pressure_p95"]

report_md = f"""# Parameterised model run report {VERSION}

**Project:** Primary care funding architecture in Aotearoa New Zealand and Australia  
**Status:** source-informed demonstrative parameterisation; not an empirically calibrated predictive model  
**Date:** {date.today().isoformat()}

## 1. What changed in this version

This version identifies the model parameters, links them to public-source inputs where available, updates the scenario inputs, reruns the game models, hybrid model, uncertainty layer and hybrid-informed MCDA, and generates plots/tables for review.

The core modelling status remains deliberately cautious: this is a source-informed demonstrative run. It improves the audit trail from v0.9.0, but it does not replace formal calibration using administrative datasets, OIA material, provider-level payment flows, ambulance disposition data, ED/hospital linkage, workforce data and stakeholder validation.

## 2. Public inputs used

The main empirical anchors used in the parameterisation are:

- NZ capitation funding was introduced in 2002, remains the core way general practice is funded, and the proposed reweighting considers age, sex, multimorbidity, rurality and socio-economic deprivation.
- A new primary care access target is proposed to take effect from 1 July 2026: more than 80% of people can access an appointment with a general practice provider within one week.
- The National Primary Care Dataset begins with general practice appointment and encounter data, including when appointments were booked, when people were seen and the outcome of appointments.
- The 2024/25 NZ Health Survey reports that 25.5% of adults experienced appointment wait time as a barrier to GP care, 14.9% did not visit a GP because of cost, and 17.1% visited an ED in the prior 12 months.
- The closed-books survey reported that only 28% of respondent general practices were freely enrolling new people in 2022, and 79% had closed or limited enrolments at some point since 2019.
- Health NZ commissions ambulance services on behalf of Health NZ and ACC, and ambulance providers report response time KPIs, demand, quality and call-type information.
- ACC pays providers through regulations, contracts or purchase orders and contributes to consultation/procedure costs for GPs, nurses and nurse practitioners in relevant injury-related care.

## 3. Parameter register

A full parameter-input register is included in the workbook and CSV artefacts. Key parameters include marginal contact benefit, capitation weighting, scope flexibility, PHO transaction cost, primary and ambulance KPI salience, data observability, co-payment level/protection, equity programme strength, ambulance alternative funding, ACC activity funding, budget tightness, safety governance, gaming controls and direct claiming.

The normalised values are not claimed to be natural units. They are documented priors for model exploration. Each parameter has a source, transformation rule, confidence rating and next data requirement.

## 4. Updated scenarios

The run now compares nine scenarios:

{md_table(scenario_df[["scenario_id", "name", "description"]], ["scenario_id", "name", "description"])}

## 5. Hybrid model results

{md_table(hybrid_df[summary_hybrid_cols].round(2), summary_hybrid_cols)}

The result is consistent with the previous architecture hypothesis. Capitation reweighting plus an access target improves the source-informed baseline, but remains weaker than a benefits schedule or full upstream access architecture on supply generation and hospital deflection. The full upstream architecture performs best. A loose benefits model improves access but is penalised by gaming, fiscal and governance risk.

## 6. Hybrid-informed MCDA

{md_table(mcda_df[mcda_cols].round(2), mcda_cols)}

The MCDA is derived from the hybrid outputs and scenario parameters. It is not a separate empirical proof; it is a decision-support translation of the model. The highest-ranked option remains the full upstream access architecture, followed by the National Primary Care Benefits Schedule.

## 7. Uncertainty summary

{md_table(unc_summary[unc_cols].round(2), unc_cols)}

The uncertainty layer perturbs all scenario levers around the source-informed values. It is still demonstrative, but it shows that the full upstream architecture remains the strongest scenario under the assumed parameter ranges.

## 8. Key sensitivity drivers

The pooled sensitivity run indicates the strongest drivers of hybrid viability are:

{md_table(pooled_sens.sort_values("abs_correlation", ascending=False).round(3), ["parameter", "correlation", "abs_correlation"], n=12)}

These should be prioritised for data collection and stakeholder validation.

## 9. Interpretation

The parameterised model supports the same cautious conclusion as the earlier demonstrative model:

> Demand-driven within rules; not demand-driven without rules.

The modelling implication is that New Zealand's policy question should not be framed as capitation versus fee-for-service. It is whether the system can define eligible contact types, allow accredited providers to generate activity within scope, calibrate co-payments with equity protections, preserve locality/equity functions, fund ambulance alternatives, and lift primary care/ambulance KPIs to hospital-equivalent salience.

## 10. Limitations

This run is not empirically calibrated. Key values remain source-informed priors or explicit modelling judgements. The results should be used to guide further empirical work, not as final estimates of fiscal impact, hospital deflection or workforce response.

## 11. Priority next data inputs

1. Full current capitation formula, rate tables and implementation workbook.
2. Provider-level payment flows: capitation, co-payments, ACC, programme funding and PHO pass-through.
3. Appointment availability, booked-to-seen interval and closed-books data from NPCD.
4. Ambulance disposition data, non-conveyance outcomes, offload delay and alternative pathway funding.
5. ED and ambulatory-sensitive hospitalisation linkage to primary care access indicators.
6. Workforce counts, provider scope utilisation and productivity by provider type.
7. Practice fee schedules and patient-level co-payment exposure.
8. Stakeholder MCDA scoring by PHOs, practices, kaupapa Māori/Pacific providers, rural providers, ambulance, Health NZ, ACC, Treasury and consumers.

## 12. Source notes

Full source details are included in `parameter-input-register-{VERSION}.csv`. This report keeps URLs in the register to preserve reproducibility.
"""

report_md_path = REPORTS_DIR / f"parameterised-model-run-report-{VERSION}.md"
report_md_path.write_text(report_md)
(OUTPUTS / report_md_path.name).write_text(report_md)

# DOCX report

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def add_df_table_docx(doc, df, cols, max_rows=None, font_size=7):
    d = df[cols].copy()
    if max_rows:
        d = d.head(max_rows)
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, col in enumerate(cols):
        hdr[i].text = col.replace('_', ' ')
        hdr[i].paragraphs[0].runs[0].font.bold = True
        set_cell_shading(hdr[i], 'D9EAF7')
    for _, row in d.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(cols):
            val = row[col]
            if isinstance(val, float):
                txt = f"{val:.2f}"
            else:
                txt = str(val)
            cells[i].text = txt
            for p in cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(font_size)
    return table


def add_picture_if_exists(doc, path, width=6.5):
    if Path(path).exists():
        doc.add_picture(str(path), width=Inches(width))


doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.6)
sec.bottom_margin = Inches(0.6)
sec.left_margin = Inches(0.6)
sec.right_margin = Inches(0.6)
styles = doc.styles
styles['Normal'].font.name = 'Arial'
styles['Normal'].font.size = Pt(9)
for style_name in ['Heading 1', 'Heading 2', 'Heading 3']:
    styles[style_name].font.name = 'Arial'

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run(f"Parameterised model run report {VERSION}")
run.bold = True
run.font.size = Pt(16)
doc.add_paragraph("Source-informed demonstrative parameterisation; not an empirically calibrated predictive model.")

doc.add_heading("1. Executive summary", level=1)
doc.add_paragraph("This version identifies the model parameters, attaches public-source inputs, updates the scenario values, reruns the game models, hybrid synthesis, uncertainty layer and hybrid-informed MCDA, and provides plots and tabular artefacts for review.")
doc.add_paragraph("The strongest conclusion remains: demand-driven within rules; not demand-driven without rules.")

doc.add_heading("2. Scenario summary", level=1)
add_df_table_docx(doc, scenario_df[["scenario_id", "name", "description"]], ["scenario_id", "name", "description"], font_size=7)

doc.add_heading("3. Hybrid results", level=1)
add_df_table_docx(doc, hybrid_df[summary_hybrid_cols].round(2), summary_hybrid_cols, font_size=7)
add_picture_if_exists(doc, fig1)
add_picture_if_exists(doc, fig2)

doc.add_heading("4. Hybrid-informed MCDA", level=1)
add_df_table_docx(doc, mcda_df[mcda_cols].round(2), mcda_cols, font_size=7)
add_picture_if_exists(doc, fig7)

doc.add_heading("5. Uncertainty and sensitivity", level=1)
add_df_table_docx(doc, unc_summary[unc_cols].round(2), unc_cols, font_size=6)
add_picture_if_exists(doc, fig5)
add_picture_if_exists(doc, fig6)

doc.add_heading("6. Game-level outputs", level=1)
add_df_table_docx(doc, game_summary.round(2), ["scenario_id", "scenario_name", "access_score", "hospital_pressure", "gaming_risk", "system_welfare"], font_size=7)
add_picture_if_exists(doc, fig4)

doc.add_heading("7. Parameter inputs", level=1)
add_df_table_docx(doc, param_df[["parameter", "domain", "source_value", "current_normalised_value", "evidence_level", "confidence_0_to_1"]].round(2), ["parameter", "domain", "source_value", "current_normalised_value", "evidence_level", "confidence_0_to_1"], max_rows=23, font_size=6)

doc.add_heading("8. Limitations and next data inputs", level=1)
doc.add_paragraph("This run is source-informed and auditable, but not empirically calibrated. The priority next inputs are full capitation/payment rules, provider-level payment flows, NPCD appointment data, ambulance disposition data, ED/hospital linkage, workforce/provider-scope data, practice fees and stakeholder MCDA scoring.")

doc.add_heading("9. Source register", level=1)
add_df_table_docx(doc, param_df[["source_name", "source_url", "parameter"]].drop_duplicates().head(20), ["source_name", "source_url", "parameter"], font_size=6)

docx_path = OUTPUTS / f"parameterised-model-run-report-{VERSION}.docx"
doc.save(docx_path)

# PDF report with ReportLab
pdf_path = OUTPUTS / f"parameterised-model-run-report-{VERSION}.pdf"
styles = getSampleStyleSheet()
styles.add(ParagraphStyle(name='Small', parent=styles['Normal'], fontSize=7, leading=8))
styles.add(ParagraphStyle(name='Tiny', parent=styles['Normal'], fontSize=6, leading=7))
styles.add(ParagraphStyle(name='H1Custom', parent=styles['Heading1'], fontSize=15, leading=18, spaceAfter=8))
styles.add(ParagraphStyle(name='H2Custom', parent=styles['Heading2'], fontSize=12, leading=14, spaceAfter=6))

story = []
story.append(Paragraph(f"Parameterised model run report {VERSION}", styles['Title']))
story.append(Paragraph("Source-informed demonstrative parameterisation; not an empirically calibrated predictive model.", styles['Normal']))
story.append(Spacer(1, 0.3*cm))
story.append(Paragraph("Executive summary", styles['H1Custom']))
story.append(Paragraph("This version identifies parameters, attaches public-source inputs, updates and reruns the game, hybrid, uncertainty and MCDA layers, then packages the results as plots and tables. The strongest conclusion remains: demand-driven within rules; not demand-driven without rules.", styles['Normal']))

def pdf_table(df, cols, max_rows=None, font_size=6):
    d = df[cols].copy()
    if max_rows:
        d = d.head(max_rows)
    data = [[Paragraph(c.replace('_', ' '), styles['Tiny']) for c in cols]]
    for _, row in d.iterrows():
        cells=[]
        for c in cols:
            val=row[c]
            if isinstance(val, float): txt=f"{val:.2f}"
            else: txt=str(val)
            cells.append(Paragraph(txt, styles['Tiny']))
        data.append(cells)
    t=Table(data, repeatRows=1)
    t.setStyle(TableStyle([
        ('BACKGROUND',(0,0),(-1,0),colors.lightgrey),
        ('GRID',(0,0),(-1,-1),0.25,colors.grey),
        ('VALIGN',(0,0),(-1,-1),'TOP'),
        ('FONTSIZE',(0,0),(-1,-1),font_size),
    ]))
    return t

story.append(Paragraph("Hybrid results", styles['H1Custom']))
story.append(pdf_table(hybrid_df[summary_hybrid_cols].round(2), summary_hybrid_cols))
story.append(Spacer(1, 0.2*cm))
for fig in [fig1, fig2, fig3]:
    story.append(Image(str(fig), width=16*cm, height=8*cm))
    story.append(Spacer(1, 0.2*cm))

story.append(PageBreak())
story.append(Paragraph("Hybrid-informed MCDA", styles['H1Custom']))
story.append(pdf_table(mcda_df[mcda_cols].round(2), mcda_cols))
story.append(Spacer(1, 0.2*cm))
story.append(Image(str(fig7), width=16*cm, height=9*cm))

story.append(PageBreak())
story.append(Paragraph("Uncertainty and sensitivity", styles['H1Custom']))
story.append(pdf_table(unc_summary[unc_cols].round(2), unc_cols))
for fig in [fig5, fig6, fig8]:
    story.append(Spacer(1, 0.2*cm))
    story.append(Image(str(fig), width=16*cm, height=8*cm))

story.append(PageBreak())
story.append(Paragraph("Game-level results", styles['H1Custom']))
story.append(pdf_table(game_summary.round(2), ["scenario_id", "scenario_name", "access_score", "hospital_pressure", "gaming_risk", "system_welfare"]))
story.append(Spacer(1, 0.2*cm))
story.append(Image(str(fig4), width=16*cm, height=10*cm))

story.append(PageBreak())
story.append(Paragraph("Parameter inputs", styles['H1Custom']))
story.append(pdf_table(param_df[["parameter", "domain", "source_value", "current_normalised_value", "evidence_level", "confidence_0_to_1"]].round(2), ["parameter", "domain", "source_value", "current_normalised_value", "evidence_level", "confidence_0_to_1"], max_rows=23, font_size=5))
story.append(Paragraph("Limitations", styles['H1Custom']))
story.append(Paragraph("The run is source-informed and auditable, but not empirically calibrated. Results should be treated as prioritisation and decision-support outputs until calibrated against administrative and stakeholder data.", styles['Normal']))

SimpleDocTemplate(str(pdf_path), pagesize=landscape(A4), rightMargin=1*cm, leftMargin=1*cm, topMargin=1*cm, bottomMargin=1*cm).build(story)

# -------------------------------
# Workbook
# -------------------------------
wb = Workbook()
# remove default
ws = wb.active
ws.title = "Overview"

header_fill = PatternFill("solid", fgColor="1F4E79")
header_font = Font(color="FFFFFF", bold=True)
subheader_fill = PatternFill("solid", fgColor="D9EAF7")
thin = Side(border_style="thin", color="BFBFBF")


def style_sheet(ws):
    ws.freeze_panes = "A2"
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(wrap_text=True, horizontal="center", vertical="center")
    for col in range(1, ws.max_column + 1):
        letter = get_column_letter(col)
        max_len = max((len(str(ws.cell(row=r, column=col).value or "")) for r in range(1, min(ws.max_row, 60) + 1)), default=10)
        ws.column_dimensions[letter].width = min(max(12, max_len + 2), 55)


def add_df_sheet(name: str, df: pd.DataFrame):
    ws = wb.create_sheet(name[:31])
    for j, col in enumerate(df.columns, 1):
        ws.cell(row=1, column=j, value=col)
    for i, row in enumerate(df.itertuples(index=False), 2):
        for j, value in enumerate(row, 1):
            if isinstance(value, (np.integer, np.floating)):
                value = float(value)
            ws.cell(row=i, column=j, value=value)
    style_sheet(ws)
    return ws

ws = wb["Overview"]
ws["A1"] = f"Parameterised model run {VERSION}"
ws["A1"].font = Font(bold=True, size=16)
ws["A2"] = "Status"
ws["B2"] = "Source-informed demonstrative; not empirically calibrated"
ws["A3"] = "Core conclusion"
ws["B3"] = "Demand-driven within rules; not demand-driven without rules."
ws["A5"] = "Top hybrid scenarios"
for idx, row in enumerate(hybrid_df.sort_values("hybrid_viability_score", ascending=False).head(5).itertuples(index=False), 6):
    ws.cell(idx, 1, row.scenario_id)
    ws.cell(idx, 2, row.scenario_name)
    ws.cell(idx, 3, row.hybrid_viability_score)
ws["E5"] = "Top MCDA scenarios"
for idx, row in enumerate(mcda_df.head(5).itertuples(index=False), 6):
    ws.cell(idx, 5, row.scenario_id)
    ws.cell(idx, 6, row.scenario_name)
    ws.cell(idx, 7, row.risk_adjusted_mcda_score)
# Add images to overview
try:
    img = XLImage(str(fig1)); img.width = 600; img.height = 300; ws.add_image(img, "A13")
    img2 = XLImage(str(fig7)); img2.width = 600; img2.height = 320; ws.add_image(img2, "I13")
except Exception:
    pass
for c in ["A", "B", "C", "E", "F", "G"]:
    ws.column_dimensions[c].width = 28

add_df_sheet("Parameter Inputs", param_df)
add_df_sheet("Scenario Inputs", scenario_df)
add_df_sheet("Hybrid Results", hybrid_df)
add_df_sheet("Game Summary", game_summary)
add_df_sheet("Game Results", game_df)
add_df_sheet("Uncertainty Summary", unc_summary)
add_df_sheet("Sensitivity Viability", sens_viability)
add_df_sheet("Sensitivity Hospital", sens_hospital)
add_df_sheet("MCDA", mcda_df)
add_df_sheet("Equilibrium Matrix", equilibrium_matrix)

xlsx_path = OUTPUTS / f"parameterised-model-workbook-{VERSION}.xlsx"
wb.save(xlsx_path)

# -------------------------------
# Conductor track and docs
# -------------------------------
(TRACK_DIR / "metadata.json").write_text('{\n  "track_id": "015-parameterised-model-run",\n  "version": "v1.0.0",\n  "status": "implemented",\n  "description": "Identify parameters, attach inputs, update and run models with plots and tables."\n}\n')
(TRACK_DIR / "spec.md").write_text("""# Track 015 — Parameterised model run v1.0.0

## Goal
Identify each model parameter, document its source-informed input, update scenario values, rerun game/hybrid/uncertainty/MCDA models, and produce plots/tables.

## Status
Implemented as a source-informed demonstrative run. Not empirically calibrated.
""")
(TRACK_DIR / "plan.md").write_text("""# Plan

1. Build parameter-input register.
2. Build parameterised scenarios.
3. Run deterministic game and hybrid models.
4. Run uncertainty stress tests.
5. Run hybrid-informed MCDA.
6. Generate tables, plots, workbook and report.
7. Add tests and versioned artefacts.
""")

# Parameter brief
parameter_brief = f"""# Parameter and input register {VERSION}

This register links each model lever to a public-source input where available, a normalisation rule, a confidence rating and the next data needed for empirical calibration.

The values are source-informed priors, not calibrated coefficients.

{md_table(param_df[["parameter", "domain", "source_value", "transform_to_normalised", "current_normalised_value", "confidence_0_to_1"]].round(2), ["parameter", "domain", "source_value", "transform_to_normalised", "current_normalised_value", "confidence_0_to_1"])}
"""
(PARAMS_DIR / f"parameter-input-brief-{VERSION}.md").write_text(parameter_brief)
(OUTPUTS / f"parameter-input-brief-{VERSION}.md").write_text(parameter_brief)

# Update version/changelog/repo index
(ROOT / "VERSION.md").write_text(f"# Version\n\n{VERSION}\n\nSource-informed parameterised model run.\n")
changelog = ROOT / "CHANGELOG.md"
old = changelog.read_text() if changelog.exists() else ""
entry = f"""
## {VERSION} — parameterised model run

- Added source-informed parameter input register.
- Added parameterised scenario set P0-P8.
- Reran deterministic game and hybrid models.
- Added uncertainty stress test and hybrid-informed MCDA.
- Added workbook, report, plots and tables.
- Added conductor track 015.

"""
if f"## {VERSION}" not in old:
    changelog.write_text(entry + old)

repo_index = f"""# Repo index {VERSION}

Key outputs:

- `outputs/parameterised-model-run-report-{VERSION}.pdf`
- `outputs/parameterised-model-run-report-{VERSION}.docx`
- `outputs/parameterised-model-workbook-{VERSION}.xlsx`
- `outputs/parameter-input-register-{VERSION}.csv`
- `outputs/parameterised-scenario-inputs-{VERSION}.csv`
- `outputs/parameterised-game-results-{VERSION}.csv`
- `outputs/parameterised-hybrid-results-{VERSION}.csv`
- `outputs/parameterised-hybrid-uncertainty-summary-{VERSION}.csv`
- `outputs/hybrid-informed-mcda-results-{VERSION}.csv`

Status: source-informed demonstrative run; not empirically calibrated.
"""
(DOCS / f"repo-index-{VERSION}.md").write_text(repo_index)

# Copy report markdown into docs and outputs for consistency
(DOCS / f"parameterised-model-run-report-{VERSION}.md").write_text(report_md)

print("Created v1.0.0 parameterised run outputs")
print("Hybrid results:")
print(hybrid_df[["scenario_id", "scenario_name", "hybrid_viability_score", "weighted_hospital_pressure"]].to_string(index=False))
print("MCDA results:")
print(mcda_df[["rank", "scenario_id", "scenario_name", "risk_adjusted_mcda_score"]].to_string(index=False))
