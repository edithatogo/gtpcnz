from pathlib import Path
import csv, textwrap, os, subprocess, sys
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT=Path('/mnt/data/work/project/primary-care-funding-architecture')
VER='v1.1.0'
OUT=ROOT/'outputs'
FIG=ROOT/'docs/figures'
OUT.mkdir(exist_ok=True)

# Helpers

def set_cell_shading(cell, fill='D9EAF7'):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), 'true')
    trPr.append(tblHeader)

def style_table(table, header_fill='1F4E79'):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8.5)
        if i == 0:
            set_repeat_table_header(row)
            for cell in row.cells:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.bold = True
                        r.font.color.rgb = None

def add_table_from_rows(doc, headers, rows, title=None, max_rows=None):
    if title:
        doc.add_paragraph(title, style='Intense Quote')
    use_rows = rows if max_rows is None else rows[:max_rows]
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i,h in enumerate(headers):
        hdr[i].text = str(h)
    for row in use_rows:
        cells = table.add_row().cells
        for i,h in enumerate(headers):
            cells[i].text = str(row.get(h,''))
    style_table(table)
    return table

def read_csv(rel):
    with (ROOT/rel).open(encoding='utf-8') as f:
        return list(csv.DictReader(f))

def add_hyperlink_like_paragraph(doc, text, url):
    p=doc.add_paragraph()
    p.add_run(text+': ').bold=True
    run=p.add_run(url)
    run.font.size = Pt(8)

# Content data
priority=read_csv(f'docs/validation/priority-empirical-checks-{VER}.csv')
threshold=read_csv(f'docs/validation/evidence-threshold-matrix-{VER}.csv')
workplan=read_csv(f'docs/validation/validation-workplan-{VER}.csv')
sources=read_csv(f'docs/validation/source-registry-{VER}.csv')
oia=read_csv(f'docs/oia/oia-data-request-tracker-{VER}.csv')
pilot=read_csv(f'docs/pilot/pilot-evaluation-design-matrix-{VER}.csv')
survey=read_csv(f'data/templates/stakeholder-game-validation-survey-{VER}.csv')

# Main report
report=Document()
sec=report.sections[0]
sec.top_margin=Inches(0.6); sec.bottom_margin=Inches(0.6); sec.left_margin=Inches(0.6); sec.right_margin=Inches(0.6)
styles=report.styles
styles['Normal'].font.name='Arial'; styles['Normal'].font.size=Pt(10)
styles['Title'].font.name='Arial'; styles['Title'].font.size=Pt(20)
for s in ['Heading 1','Heading 2','Heading 3']:
    styles[s].font.name='Arial'

p=report.add_paragraph()
p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Primary care funding architecture: pragmatic validation package')
r.bold=True; r.font.size=Pt(18)
p2=report.add_paragraph()
p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
p2.add_run(f'{VER} - 8 May 2026').italic=True
report.add_paragraph('Purpose: proceed with stakeholder validation, rapid review, OIA/data requests, targeted empirical checks and pilot design, without attempting a fully calibrated predictive model at this stage.')

report.add_heading('1. Executive summary', level=1)
for para in [
    'The project should now proceed as a validation and translation programme rather than a full national predictive modelling programme. The current model is source-informed and parameterised, but it remains demonstrative. That is sufficient for policy scoping, RACMA discussion, stakeholder engagement, Substack publication and an NZMJ Viewpoint or methods/protocol article.',
    'A fully calibrated predictive model becomes necessary only if decision-makers require quantified forecasts of ED presentations, hospital admissions, costs, savings, supply effects or implementation business-case outcomes.',
    'The immediate work should therefore focus on validating the games, exposing disagreements, making funding flows observable and testing the five most load-bearing assumptions.'
]:
    report.add_paragraph(para)

report.add_heading('2. Evidence threshold by output', level=1)
add_table_from_rows(report, ['output','minimum_evidence_needed','current_status','do_not_claim'], threshold, max_rows=None)
report.add_picture(str(FIG/f'evidence-thresholds-{VER}.png'), width=Inches(6.7))

report.add_heading('3. Priority empirical checks', level=1)
report.add_paragraph('These five assumptions carry the greatest decision weight. They should be tested before a national implementation business case, but they do not need to be fully resolved before public or professional policy discussion begins.')
add_table_from_rows(report, ['priority_rank','parameter','why_it_matters','proposed_method','decision_use'], priority, max_rows=None)
report.add_picture(str(FIG/f'priority-empirical-checks-{VER}.png'), width=Inches(6.7))

report.add_heading('4. Pragmatic validation pathway', level=1)
add_table_from_rows(report, ['phase','timeframe','workstream','activities','outputs'], workplan, max_rows=None)
report.add_picture(str(FIG/f'validation-roadmap-{VER}.png'), width=Inches(6.7))

report.add_heading('5. Stakeholder validation and MCDA', level=1)
report.add_paragraph('The highest-value next step is stakeholder validation. The game-informed MCDA should not be treated as proof. It is a deliberative decision-support method that makes assumptions, values, disagreement and risk tolerance explicit.')
report.add_paragraph('Suggested stakeholder groups include general practice, nurse practitioners, nurses, pharmacists, physiotherapists, mental health providers, paramedics, rural providers, kaupapa Maori and Pacific providers, PHOs/localities, Health NZ, ACC, ambulance leaders, Treasury/fiscal observers and consumer advocates.')
report.add_paragraph('Each game should be scored for reality, harm, hospital-growth contribution, equity relevance, tractability, reform risk and confidence. Policy options should then be scored against access/supply, hospital deflection, equity, rural resilience, fiscal control, gaming risk, market entry, clinical governance, political feasibility and data readiness.')

report.add_heading('6. OIA and data requests', level=1)
add_table_from_rows(report, ['request_id','agency','request_title','purpose','status'], oia, max_rows=None)
report.add_paragraph('The OIA pack in this release contains draft request text for each item. The priority is not to obtain every dataset immediately, but to establish which funding and data flows are observable, which are withheld, and which would require formal research agreements.')

report.add_heading('7. Rapid scoping review', level=1)
report.add_paragraph('The rapid review should be PRISMA-ScR aligned and should map mechanisms rather than only outcomes. The key question is: what evidence exists that funding architecture affects access, marginal supply, market entry, equity, ambulance/urgent-care pathways and hospital demand?')
report.add_paragraph('The review should produce an evidence map linked to the 14 games and the five priority assumptions.')

report.add_heading('8. Prospective pilot/evaluation design', level=1)
add_table_from_rows(report, ['pilot_domain','pilot_question','primary_outcomes','design'], pilot, max_rows=None)
report.add_paragraph('The pilot pathway should test contact types, provider scope, ambulance alternatives, direct claiming and equity protections before any national-scale implementation. A stepped-wedge or matched comparison design is preferable where feasible.')

report.add_heading('9. Recommended language', level=1)
report.add_paragraph('Preferred formulation: The current work is a structured, source-informed and falsifiable policy hypothesis. It is sufficient to open the policy conversation and design validation. It is not yet a calibrated estimate of effect size.', style=None)
report.add_paragraph('Avoid claiming: quantified ED reduction, fiscal savings, proven PHO causation, or definitive market failure without targeted empirical evidence.')

report.add_heading('10. Source anchors', level=1)
for s in sources:
    add_hyperlink_like_paragraph(report, s['source_id']+' - '+s['name'], s['url'])

# Footer
for section in report.sections:
    footer = section.footer.paragraphs[0]
    footer.text = f'Primary Care Funding Architecture {VER} | pragmatic validation without full predictive calibration'
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

report_path=OUT/f'validation-and-translation-report-{VER}.docx'
report.save(report_path)

# Workshop pack docx
wk=Document()
sec=wk.sections[0]
sec.top_margin=Inches(0.6); sec.bottom_margin=Inches(0.6); sec.left_margin=Inches(0.6); sec.right_margin=Inches(0.6)
wk.styles['Normal'].font.name='Arial'; wk.styles['Normal'].font.size=Pt(10)
for s in ['Heading 1','Heading 2','Heading 3']:
    wk.styles[s].font.name='Arial'

p=wk.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Stakeholder game-validation and MCDA workshop pack')
r.bold=True; r.font.size=Pt(18)
p=wk.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run(f'{VER} - 8 May 2026').italic=True
wk.add_paragraph('Purpose: test whether the mapped policy games are real, important, tractable and risky; then compare policy options with explicit criteria and weights.')

wk.add_heading('1. Pre-work for participants', level=1)
for item in ['Read the one-page summary of the 14 games.', 'Bring one example from your own work where upstream access, funding rules or accountability affected patient flow.', 'Be prepared to score both mechanisms and policy options. Disagreement is useful and expected.']:
    wk.add_paragraph(item, style='List Bullet')

wk.add_heading('2. Agenda', level=1)
agendas=[
    {'Time':'0:00-0:15','Item':'Context and caveats','Output':'Shared understanding that this is validation, not predictive proof.'},
    {'Time':'0:15-0:45','Item':'Game-position scoring','Output':'Scores for 14 games.'},
    {'Time':'0:45-1:10','Item':'Breakout discussion','Output':'Objections, missing games and evidence needs.'},
    {'Time':'1:10-1:45','Item':'Policy-option MCDA','Output':'Scored policy options under common criteria.'},
    {'Time':'1:45-2:00','Item':'Priority assumptions and next steps','Output':'Top empirical checks and pilot priorities.'},
]
add_table_from_rows(wk, ['Time','Item','Output'], agendas)

wk.add_heading('3. Game validation template', level=1)
add_table_from_rows(wk, ['game_id','game_name','mechanism_to_validate','is_this_game_real_1_5','harm_if_unresolved_1_5','hospital_growth_contribution_1_5','confidence_1_5'], survey, max_rows=None)

wk.add_heading('4. MCDA criteria', level=1)
crit_rows=[{'criterion_id':cid,'criterion_name':name,'criterion_description':desc} for cid,name,desc in [
    ('C1','Access and supply generation','Does the option increase safe upstream capacity?'),
    ('C2','Hospital deflection','Does it reduce avoidable ED, ambulance and hospital flow?'),
    ('C3','Equity and Te Tiriti legitimacy','Does access improve without worsening inequity?'),
    ('C4','Rural and in-person resilience','Does it protect local care, not just telehealth?'),
    ('C5','Fiscal sustainability','Is the model affordable and controllable?'),
    ('C6','Gaming and low-value activity risk','Does it avoid opportunistic or low-value activity?'),
    ('C7','Administrative simplicity and market entry','Does it lower transaction costs and entry barriers?'),
    ('C8','Governance and clinical safety','Are scope, prescribing, audit and safety controls strong?'),
    ('C9','Political feasibility','Can the option survive contestation?'),
    ('C10','Data and accountability readiness','Can outcomes be measured and managed?')]]
add_table_from_rows(wk, ['criterion_id','criterion_name','criterion_description'], crit_rows)

wk.add_heading('5. Scoring rules', level=1)
score_rows=[
    {'Score':'-2','Meaning':'Substantially worsens criterion'},
    {'Score':'-1','Meaning':'Slightly worsens criterion'},
    {'Score':'0','Meaning':'No meaningful change'},
    {'Score':'+1','Meaning':'Slightly improves criterion'},
    {'Score':'+2','Meaning':'Substantially improves criterion'},
]
add_table_from_rows(wk, ['Score','Meaning'], score_rows)

wk.add_heading('6. Expected outputs', level=1)
for item in ['Group-level game scores and disagreement map.', 'Weighted and unweighted policy-option ranking.', 'List of assumptions requiring empirical testing.', 'Suggested pilot domains and implementation risks.', 'Revisions to the game map and policy brief.']:
    wk.add_paragraph(item, style='List Bullet')

for section in wk.sections:
    footer = section.footer.paragraphs[0]
    footer.text = f'Primary Care Funding Architecture {VER} | stakeholder workshop pack'
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

wk_path=OUT/f'stakeholder-mcda-workshop-pack-{VER}.docx'
wk.save(wk_path)
print(report_path)
print(wk_path)
